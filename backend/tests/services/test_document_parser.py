import fitz
import pytest

from docx import Document

from app.schemas.document import DocumentType
from app.services.document_parser import (
    DocumentParser,
    UnsupportedDocumentTypeError,
)


def test_parse_txt_document(tmp_path):
    file = tmp_path / "curriculo.txt"

    file.write_text(
        """
        João Silva

        Senior Software Engineer

        Java
        AWS
        Kafka
        """,
        encoding="utf-8",
    )

    parser = DocumentParser()

    document = parser.parse(file)

    assert document.filename == "curriculo.txt"
    assert document.document_type == DocumentType.TXT
    assert document.content == (
        "João Silva\n"
        "Senior Software Engineer\n"
        "Java\n"
        "AWS\n"
        "Kafka"
    )
    assert document.character_count > 0

    assert document.page_count == 1
    assert document.pages[0].number == 1
    assert "João Silva" in document.pages[0].content


def test_document_not_found():
    parser = DocumentParser()

    with pytest.raises(FileNotFoundError):
        parser.parse("arquivo_inexistente.txt")


def test_unsupported_document_type(tmp_path):
    file = tmp_path / "curriculo.xyz"
    file.write_text("teste", encoding="utf-8")

    parser = DocumentParser()

    with pytest.raises(UnsupportedDocumentTypeError):
        parser.parse(file)


def test_parse_pdf_document(tmp_path):
    file = tmp_path / "curriculo.pdf"

    pdf = fitz.open()

    page = pdf.new_page()

    page.insert_text(
        (72, 72),
        "João Silva\nSenior Software Engineer\nJava AWS Kafka",
    )

    pdf.save(file)
    pdf.close()

    parser = DocumentParser()

    document = parser.parse(file)

    assert document.filename == "curriculo.pdf"
    assert document.document_type == DocumentType.PDF

    assert "João Silva" in document.content
    assert "Senior Software Engineer" in document.content
    assert "Kafka" in document.content

    assert document.character_count > 0
    assert document.page_count == 1

    assert document.pages[0].number == 1
    assert "João Silva" in document.pages[0].content


def test_parse_multi_page_pdf(tmp_path):
    file = tmp_path / "curriculo_multiplas_paginas.pdf"

    pdf = fitz.open()

    page_one = pdf.new_page()
    page_one.insert_text(
        (72, 72),
        "João Silva\nSenior Software Engineer",
    )

    page_two = pdf.new_page()
    page_two.insert_text(
        (72, 72),
        "Java\nAWS\nKafka",
    )

    pdf.save(file)
    pdf.close()

    parser = DocumentParser()

    document = parser.parse(file)

    assert document.document_type == DocumentType.PDF
    assert document.page_count == 2

    assert document.pages[0].number == 1
    assert document.pages[1].number == 2

    assert "João Silva" in document.pages[0].content
    assert "Kafka" in document.pages[1].content

    assert "João Silva" in document.content
    assert "Kafka" in document.content


def test_parse_docx_document(tmp_path):
    file = tmp_path / "curriculo.docx"

    doc = Document()
    doc.add_paragraph("João Silva")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("Java")
    doc.add_paragraph("AWS")
    doc.add_paragraph("Kafka")
    doc.save(file)

    parser = DocumentParser()

    document = parser.parse(file)

    assert document.filename == "curriculo.docx"
    assert document.document_type == DocumentType.DOCX

    assert "João Silva" in document.content
    assert "Senior Software Engineer" in document.content
    assert "Kafka" in document.content

    assert document.character_count > 0
    assert document.page_count == 1

    assert document.pages[0].number == 1
    assert "João Silva" in document.pages[0].content

def test_parse_raw_text():
    parser = DocumentParser()

    document = parser.parse_text(
        """
        Senior Software Engineer

        Buscamos profissional com experiência em Java,
        AWS e Kafka.

        Liderança técnica será um diferencial.
        """,
        source_name="vaga_copiada",
    )

    assert document.filename == "vaga_copiada"
    assert document.document_type == DocumentType.RAW_TEXT

    assert "Senior Software Engineer" in document.content
    assert "Java" in document.content
    assert "AWS" in document.content
    assert "Kafka" in document.content

    assert document.character_count > 0
    assert document.page_count == 1

    assert document.pages[0].number == 1
    assert "Liderança técnica" in document.pages[0].content

def test_parse_raw_text_empty():
    parser = DocumentParser()

    with pytest.raises(ValueError):
        parser.parse_text("   ")