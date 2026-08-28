import pymupdf

from pathlib import Path
from docx import Document

from app.schemas.document import (
    DocumentPage,
    DocumentType,
    ParsedDocument,
)

class UnsupportedDocumentTypeError(Exception):
    pass


class DocumentParser:
    SUPPORTED_TYPES = {
        ".txt": DocumentType.TXT,
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
    }

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Arquivo não encontrado: {path}"
            )

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_TYPES:
            raise UnsupportedDocumentTypeError(
                f"Formato não suportado: {extension}"
            )

        document_type = self.SUPPORTED_TYPES[extension]

        if document_type == DocumentType.TXT:
            pages = self._parse_txt(path)

        elif document_type == DocumentType.PDF:
            pages = self._parse_pdf(path)

        elif document_type == DocumentType.DOCX:
            pages = self._parse_docx(path)

        else:
            raise NotImplementedError(
                f"Parser para {document_type.value} ainda não implementado."
            )

        normalized_pages = [
            DocumentPage(
                number=page.number,
                content=self._normalize_text(page.content),
            )
            for page in pages
        ]

        content = "\n".join(
            page.content
            for page in normalized_pages
            if page.content
        )

        return ParsedDocument(
            filename=path.name,
            document_type=document_type,
            content=content,
            character_count=len(content),
            page_count=len(normalized_pages),
            pages=normalized_pages,
        )

    def _normalize_text(self, text: str) -> str:
        lines = [
            line.strip()
            for line in text.splitlines()
        ]

        lines = [
            line
            for line in lines
            if line
        ]

        return "\n".join(lines)

    def parse_text(
        self,
        text: str,
        source_name: str = "texto_colado",
    ) -> ParsedDocument:
        if not text or not text.strip():
            raise ValueError("O texto informado não pode estar vazio.")

        normalized_content = self._normalize_text(text)

        page = DocumentPage(
            number=1,
            content=normalized_content,
        )

        return ParsedDocument(
            filename=source_name,
            document_type=DocumentType.RAW_TEXT,
            content=normalized_content,
            character_count=len(normalized_content),
            page_count=1,
            pages=[page],
        )

    def _parse_txt(self, path: Path) -> list[DocumentPage]:
        content = path.read_text(encoding="utf-8")

        return [
            DocumentPage(
                number=1,
                content=content,
            )
        ]

    def _parse_pdf(self, path: Path) -> list[DocumentPage]:
        pages = []

        with pymupdf.open(path) as document:
            for index, page in enumerate(document):
                pages.append(
                    DocumentPage(
                        number=index + 1,
                        content=page.get_text(),
                    )
                )

        return pages

    def _parse_docx(self, path: Path) -> list[DocumentPage]:
        document = Document(path)

        content = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )

        return [
            DocumentPage(
                number=1,
                content=content,
            )
        ]